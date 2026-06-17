from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_table(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    for row in data:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

def main():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    
    section = doc.sections[0]
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1.5)
    
    add_heading(doc, '编译原理课程设计报告', level=1)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('课程名称：编译原理\n\n').bold = True
    
    add_heading(doc, '一、课程设计目的', level=2)
    add_paragraph(doc, '本课程设计旨在实现一个完整的编译器前端系统，包括词法分析器、语法分析器和语义分析器，深入理解编译原理的核心概念和实现技术。')
    
    add_heading(doc, '1.1 设计目标', level=3)
    doc.add_paragraph('• 实现词法分析器：对源程序进行扫描，识别并分类各类单词符号')
    doc.add_paragraph('• 实现语法分析器：基于递归下降分析法，验证程序的语法正确性')
    doc.add_paragraph('• 实现语义分析器：构建符号表，生成四元式中间代码')
    
    add_heading(doc, '二、系统概述', level=2)
    
    add_heading(doc, '2.1 系统架构', level=3)
    add_paragraph(doc, '编译器前端系统由三个主要模块组成：词法分析器、语法分析器和语义分析器。')
    
    add_code_block(doc, '┌─────────────────────────────────────────────────────────────┐\n'
                       '│                    编译器前端系统                           │\n'
                       '├─────────────────────────────────────────────────────────────┤\n'
                       '│  ┌─────────────┐    ┌─────────────┐    ┌───────────────┐   │\n'
                       '│  │  词法分析器  │───▶│  语法分析器  │───▶│  语义分析器   │   │\n'
                       '│  │ TokenScanner│    │SyntaxAnalyzer│    │SemanticAnalyzer│   │\n'
                       '│  └─────────────┘    └─────────────┘    └───────────────┘   │\n'
                       '│        │                  │                  │              │\n'
                       '│        ▼                  ▼                  ▼              │\n'
                       '│   输入源程序          语法树/产生式        四元式/符号表    │\n'
                       '└─────────────────────────────────────────────────────────────┘')
    
    add_heading(doc, '2.2 文件结构', level=3)
    file_data = [
        ['Tokenscaner.h', '词法分析器头文件', '定义Token类型、符号表、TokenAnalyzer类'],
        ['Tokenscaner.cpp', '词法分析器实现', '实现词法扫描、单词识别功能'],
        ['SyntaxAnalyzer.h', '语法分析器头文件', '定义语法分析器类和产生式规则'],
        ['SyntaxAnalyzer.cpp', '语法分析器实现', '实现递归下降语法分析'],
        ['SemanticAnalyzer.h', '语义分析器头文件', '定义符号表、四元式结构'],
        ['SemanticAnalyzer.cpp', '语义分析器实现', '实现符号表构建和四元式生成']
    ]
    add_table(doc, file_data, ['文件', '功能', '说明'])
    
    add_heading(doc, '三、词法分析器设计与实现', level=2)
    
    add_heading(doc, '3.1 设计思想', level=3)
    add_paragraph(doc, '词法分析器负责将源程序字符流转换为单词符号序列。采用状态机思想，通过逐个字符扫描，根据字符类型进入不同的处理分支。')
    
    add_heading(doc, '3.2 核心数据结构', level=3)
    
    add_heading(doc, '3.2.1 Token类型定义', level=4)
    add_code_block(doc, 'enum class TokenType\n{\n    KEYWORD,        // 关键字：if, else, while, return, int, float\n    IDENTIFIED,     // 标识符\n    NUMBER,         // 数字\n    STRING,         // 字符串\n    OPERATOR,       // 运算符\n    DELIMITED,      // 分隔符\n    END             // 文件结束标记\n};')
    
    add_heading(doc, '3.2.2 Token结构', level=4)
    add_code_block(doc, 'struct Token\n{\n    TokenType type;      // 单词类型\n    string lexeme;       // 单词值\n    int line, column;    // 位置信息\n    int int_value;       // 整数值\n    double double_value; // 浮点数值\n    string string_value; // 字符串值\n};')
    
    add_heading(doc, '3.3 主要功能实现', level=3)
    
    add_heading(doc, '3.3.1 关键字识别', level=4)
    add_code_block(doc, 'inline unordered_map<string, TokenType> keywordTable = {\n    {"if", TokenType::KEYWORD},\n    {"else", TokenType::KEYWORD},\n    {"while", TokenType::KEYWORD},\n    {"return", TokenType::KEYWORD},\n    {"int", TokenType::KEYWORD},\n    {"float", TokenType::KEYWORD}\n};')
    
    add_heading(doc, '3.4 测试结果', level=3)
    add_paragraph(doc, '输入：', bold=True)
    add_code_block(doc, 'while (i>=j) i--;')
    add_paragraph(doc, '输出：', bold=True)
    add_code_block(doc, '< while, keyword >\n< (, delim >\n< i, id >\n< >=, op >\n< j, id >\n< ), delim >\n< i, id >\n< --, op >\n< ;, delim >')
    
    add_heading(doc, '四、语法分析器设计与实现', level=2)
    
    add_heading(doc, '4.1 设计思想', level=3)
    add_paragraph(doc, '采用递归下降分析法，为每个非终结符编写一个递归函数。语法分析器基于LL(1)文法，通过向前看一个token来决定使用哪个产生式。')
    
    add_heading(doc, '4.2 文法定义', level=3)
    add_code_block(doc, 'stmts → stmt rest0\n'
                       'rest0 → stmt rest0 | ε\n'
                       '\n'
                       'stmt → loc = expr ;\n'
                       '     | if ( bool ) stmt [ else stmt ]\n'
                       '     | while ( bool ) stmt\n'
                       '\n'
                       'loc → id resta\n'
                       'resta → [ elist ] | ε\n'
                       '\n'
                       'expr → term rest5\n'
                       'rest5 → + term rest5 | - term rest5 | ε\n'
                       '\n'
                       'term → unary rest6\n'
                       'rest6 → * unary rest6 | / unary rest6 | ε')
    
    add_heading(doc, '4.3 递归下降分析器实现', level=3)
    
    func_data = [
        ['stmts()', 'stmts → stmt rest0', '语句序列'],
        ['stmt()', 'stmt → loc=expr; | if... | while...', '单条语句'],
        ['loc()', 'loc → id resta', '左值（变量或数组元素）'],
        ['expr()', 'expr → term rest5', '表达式'],
        ['term()', 'term → unary rest6', '项'],
        ['factor()', 'factor → (expr) | loc | num', '因子']
    ]
    add_table(doc, func_data, ['函数', '对应产生式', '功能'])
    
    add_heading(doc, '4.4 测试结果', level=3)
    add_paragraph(doc, '输入：', bold=True)
    add_code_block(doc, 'while(sum<10000)\n'
                       'if(a<b)\n'
                       'sum=sum*(c[10]+10);\n'
                       'else\n'
                       'c[10]=sum*c[10]+10;\n'
                       'x[i,j]=sum;')
    add_paragraph(doc, '输出（部分产生式）：', bold=True)
    add_code_block(doc, '(1)stmts -> stmt rest0\n'
                       '(2)stmt -> while ( bool ) stmt\n'
                       '(3)bool -> equality\n'
                       '(4)equality -> rel rest4\n'
                       '(5)rel -> expr rop_expr\n'
                       '(6)expr -> term rest5\n'
                       '(7)term -> unary rest6\n'
                       '(8)unary -> factor\n'
                       '(9)factor -> loc\n'
                       '(10)loc -> id resta\n'
                       '(11)resta -> ε\n'
                       '...')
    
    add_heading(doc, '五、语义分析器设计与实现', level=2)
    
    add_heading(doc, '5.1 设计思想', level=3)
    add_paragraph(doc, '语义分析器在语法分析的基础上，完成以下功能：\n1. 构建符号表，记录变量的名称、类型、作用域等信息\n2. 生成四元式中间代码，表示程序的语义动作')
    
    add_heading(doc, '5.2 核心数据结构', level=3)
    
    add_heading(doc, '5.2.1 四元式结构', level=4)
    add_code_block(doc, 'struct Quadruple {\n    string op;     // 操作符\n    string arg1;   // 操作数1\n    string arg2;   // 操作数2\n    string result; // 结果\n    int label;     // 用于回填的标记\n};')
    
    add_heading(doc, '5.3 主要功能实现', level=3)
    
    add_heading(doc, '5.3.1 符号表管理', level=4)
    add_code_block(doc, 'void insertSymbol(const string &name, const TokenType &type, int line) {\n    if (lookupSymbol(name) == nullptr) {\n        scanner.table.insert(name, SymbolEntry(name, type, currentScope, line));\n    }\n}\n\nSymbolEntry* lookupSymbol(const string &name) {\n    return scanner.table.lookup(name);\n}')
    
    add_heading(doc, '5.3.2 四元式生成', level=4)
    add_code_block(doc, 'void emit(const string &op, const string &arg1, const string &arg2, const string &result) {\n    quadList.emit(op, arg1, arg2, result);\n}\n\nstring newtemp() {\n    return "t" + to_string(tempCounter++);\n}')
    
    add_heading(doc, '5.3.3 回填机制', level=4)
    add_code_block(doc, 'void backpatch(int list, int target) {\n    while (list != -1) {\n        int next = quadList[list].label;\n        quadList[list].result = to_string(target);\n        list = next;\n    }\n}\n\nint merge(int list1, int list2) {\n    if (list1 == -1) return list2;\n    if (list2 == -1) return list1;\n    quadList[list1].label = list2;\n    return list1;\n}')
    
    add_heading(doc, '5.4 表达式翻译', level=3)
    
    add_heading(doc, '5.4.1 算术表达式', level=4)
    add_code_block(doc, 'string expr() {\n    string termPlace = term();\n    return rest5(termPlace);\n}\n\nstring rest5(string inPlace) {\n    if (currentToken.lexeme == "+") {\n        currentToken = scanner.scan();\n        string termPlace = term();\n        string t = newtemp();\n        emit("+", inPlace, termPlace, t);\n        return rest5(t);\n    }\n    return inPlace;\n}')
    
    add_heading(doc, '5.5 控制流语句翻译', level=3)
    
    add_heading(doc, '5.5.1 while循环', level=4)
    add_code_block(doc, 'else if (currentToken.lexeme == "while") {\n    currentToken = scanner.scan();\n    int m1 = nextquad();           // 循环开始位置\n    expect("(");\n    boolExprNode boolNode = bool_expr();\n    expect(")");\n    int m2 = nextquad();           // 循环体开始位置\n    controlNode stmt1 = stmt();\n    backpatch(stmt1.nextlist, m1); // 循环体结束跳回判断\n    backpatch(boolNode.truelist, m2); // 条件为真进入循环体\n    ret.nextlist = boolNode.falselist;\n    emit("j", "-", "-", to_string(m1)); // 无条件跳回判断\n}')
    
    add_heading(doc, '5.5.2 if-else语句', level=4)
    add_code_block(doc, 'else if (currentToken.lexeme == "if") {\n    currentToken = scanner.scan();\n    expect("(");\n    boolExprNode boolNode = bool_expr();\n    expect(")");\n    int m1 = nextquad();\n    controlNode stmt1 = stmt();\n    if (currentToken.lexeme == "else") {\n        currentToken = scanner.scan();\n        int m2 = nextquad();\n        backpatch(boolNode.truelist, m1);\n        int temp = merge(boolNode.falselist, stmt1.nextlist);\n        controlNode stmt2 = stmt();\n        ret.nextlist = merge(stmt2.nextlist, temp);\n        backpatch(boolNode.falselist, m2);\n    }\n}')
    
    add_heading(doc, '5.6 测试结果', level=3)
    add_paragraph(doc, '输入：', bold=True)
    add_code_block(doc, 'while(sum<10000)\n'
                       'if(a<b)\n'
                       'sum=sum*(c[10]+10);\n'
                       'else\n'
                       'c[10]=sum*c[10]+10;\n'
                       'x[i,j]=sum;')
    add_paragraph(doc, '输出（四元式）：', bold=True)
    add_code_block(doc, '0:(j<, sum, 10000, 2)\n'
                       '1:(j, -, -, 19)\n'
                       '2:(j<, a, b, 4)\n'
                       '3:(j, -, -, 0)\n'
                       '4:(-, c, C, t0)\n'
                       '5:(*, w, 10, t1)\n'
                       '6:(=[],, t0, t1, t2)\n'
                       '7:(+, t2, 10, t3)\n'
                       '8:(*, sum, t3, t4)\n'
                       '9:(=, t4, -, sum)')
    
    add_heading(doc, '六、测试结果汇总', level=2)
    
    add_heading(doc, '6.1 测试用例', level=3)
    test_data = [
        ['arrayTest.txt', '多维数组访问 x=A[i,j];', '通过'],
        ['boolTest.txt', '嵌套控制流（if-while-if-else）', '通过'],
        ['calculateTest.txt', '算术表达式 a=6/b+5*c-d;', '通过'],
        ['controlTest.txt', 'while循环+if-else', '通过']
    ]
    add_table(doc, test_data, ['测试文件', '测试内容', '状态'])
    
    add_heading(doc, '6.2 测试结果详情', level=3)
    
    add_heading(doc, '6.2.1 数组测试', level=4)
    add_paragraph(doc, '输入：x=A[i, j];', bold=True)
    add_code_block(doc, '0:(*, i, 2, t0)\n'
                       '1:(+, t0, j, t0)\n'
                       '2:(-, A, C, t1)\n'
                       '3:(*, w, t0, t2)\n'
                       '4:(=[],, t1, t2, t3)\n'
                       '5:(=, t3, -, x)')
    
    add_heading(doc, '6.2.2 控制流测试', level=4)
    add_paragraph(doc, '输入：', bold=True)
    add_code_block(doc, 'while(a<b)\n'
                       'if(c)\n'
                       'x=y+z;\n'
                       'else\n'
                       'x=y-z;\n'
                       'a=y;')
    add_paragraph(doc, '四元式：', bold=True)
    add_code_block(doc, '0:(j<, a, b, 2)\n'
                       '1:(j, -, -, 9)\n'
                       '2:(jnz, c, -, 4)\n'
                       '3:(j, -, -, 0)\n'
                       '4:(+, y, z, t0)\n'
                       '5:(=, t0, -, x)\n'
                       '6:(-, y, z, t1)\n'
                       '7:(=, t1, -, x)\n'
                       '8:(j, -, -, 0)\n'
                       '9:(=, y, -, a)')
    
    add_heading(doc, '七、总结', level=2)
    
    add_heading(doc, '7.1 完成的功能', level=3)
    add_paragraph(doc, '1. **词法分析器**：成功实现了单词扫描、关键字识别、运算符识别、数字识别等功能')
    add_paragraph(doc, '2. **语法分析器**：实现了递归下降分析，支持赋值语句、if-else条件语句、while循环、数组访问、表达式运算')
    add_paragraph(doc, '3. **语义分析器**：实现了符号表构建、四元式生成、控制流回填等功能')
    
    add_heading(doc, '7.2 技术亮点', level=3)
    doc.add_paragraph('• 采用LL(1)文法设计，确保语法分析的确定性')
    doc.add_paragraph('• 实现了完整的控制流语句翻译（if-else、while）')
    doc.add_paragraph('• 支持多维数组的语义分析')
    doc.add_paragraph('• 采用回填技术处理控制流跳转')
    doc.add_paragraph('• 符号表支持作用域管理')
    
    add_heading(doc, '7.3 待改进之处', level=3)
    doc.add_paragraph('• 可增加类型检查功能')
    doc.add_paragraph('• 可扩展支持更多数据类型')
    doc.add_paragraph('• 可优化四元式生成算法')
    doc.add_paragraph('• 可增加错误处理和提示信息')
    
    add_heading(doc, '八、参考文献', level=2)
    doc.add_paragraph('1. 《编译原理》（第2版），陈火旺等著')
    doc.add_paragraph('2. 《Compilers: Principles, Techniques, and Tools》（第2版），Aho, Sethi, Ullman著')
    doc.add_paragraph('3. 相关课程讲义和实验指导')
    
    add_paragraph(doc, '\n\n课程报告完成时间：2026年5月\n\n作者：[姓名]\n\n学号：[学号]\n\n班级：[班级]')
    
    doc.save('C:\\Users\\10497\\Documents\\Compliers\\课程报告_完整版.docx')
    print("Word文档已成功生成！")

if __name__ == '__main__':
    main()
